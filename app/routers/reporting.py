import os
import tempfile
import asyncio
import time
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, Query, Header
from fastapi.responses import FileResponse
from sqlalchemy import and_, asc
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from PIL import Image
from docx.shared import Inches, Pt

# Local Imports
from models import DailyReport, TwoWeeklyReport, CommonReport
from database import get_dept_context, get_dept_resources
from utils import _attach_images, sync_json_backup, background_search_logic

import io
import tempfile
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi import APIRouter, Header, Query, HTTPException
from fastapi.responses import FileResponse



from world import SSE_SESSIONS, SEARCH_TASKS 

router = APIRouter(
    prefix="/reports",
    tags=["Reporting"]
)


#================================================================================================
@router.post("/save-daily/{dateStr}")
def save_daily_report(
    dateStr: str,
    payload: list[dict],
    ctx = Depends(get_dept_context)
):
#------------------------------------------------------------------------------------------------
    db, res = ctx
    incoming_uuids = [item.get("uuid") for item in payload if item.get("uuid")]

    db.query(DailyReport).filter(
        DailyReport.report_dateStr == dateStr,
        ~DailyReport.uuid.in_(incoming_uuids)
    ).delete(synchronize_session=False)

    for index, item in enumerate(payload):
        uuid_val = item.get("uuid")
        report = db.query(DailyReport).filter(DailyReport.uuid == uuid_val).first()

        if not report:
            report = DailyReport(uuid=uuid_val, report_dateStr=dateStr)
            db.add(report)
        
        report.text_content = item.get("text_content")
        report.sort_order = index
        report.images = [] 
        _attach_images(db, report, item.get("image_uuids", []))

    db.commit()
    # Trigger Backup
    sync_json_backup(db, res, "daily", dateStr)
    return True



#================================================================================================
@router.post("/save-2weekly/{rangeStr}")
def save_2weekly_report(
    rangeStr: str,
    payload: list[dict],
    ctx = Depends(get_dept_context)
):
#------------------------------------------------------------------------------------------------
    db, res = ctx
    incoming_uuids = [item.get("uuid") for item in payload if item.get("uuid")]

    db.query(TwoWeeklyReport).filter(
        TwoWeeklyReport.report_rangeStr == rangeStr,
        ~TwoWeeklyReport.uuid.in_(incoming_uuids)
    ).delete(synchronize_session=False)

    for index, item in enumerate(payload):
        uuid_val = item.get("uuid")
        report = db.query(TwoWeeklyReport).filter(TwoWeeklyReport.uuid == uuid_val).first()

        if not report:
            report = TwoWeeklyReport(uuid=uuid_val, report_rangeStr=rangeStr)
            db.add(report)

        report.text_content = item.get("text_content")
        report.sort_order = index
        report.images = []
        _attach_images(db, report, item.get("image_uuids", []))

    db.commit()
    # Trigger Backup
    sync_json_backup(db, res, "2-weekly", rangeStr)
    return True



#==================================================================================================
@router.post("/save-common")
def save_common_report(
    payload: list[dict],
    ctx = Depends(get_dept_context)
):
#--------------------------------------------------------------------------------------------------
    db, res = ctx
    incoming_uuids = [item.get("uuid") for item in payload if item.get("uuid")]

    db.query(CommonReport).filter(~CommonReport.uuid.in_(incoming_uuids)).delete(synchronize_session=False)

    for index, item in enumerate(payload):
        uuid_val = item.get("uuid")
        report = db.query(CommonReport).filter(CommonReport.uuid == uuid_val).first()

        if not report:
            report = CommonReport(uuid=uuid_val)
            db.add(report)

        report.text_content = item.get("text_content")
        report.sort_order = index

    db.commit()
    return True



#==================================================================================================
@router.get("/load-daily/{dateStr}")
def load_daily_report(
    dateStr: str,
    ctx = Depends(get_dept_context)
):
#--------------------------------------------------------------------------------------------------
    db, _ = ctx

    reports = db.query(DailyReport).filter(
        DailyReport.report_dateStr == dateStr
    ).order_by(asc(DailyReport.sort_order)).all()

    print(len(reports))
    return [
        {
            "uuid": r.uuid,
            "text_content": r.text_content,
            "image_uuids": [img.uuid for img in r.images]
        }
        for r in reports
    ]



#==================================================================================================
@router.get("/load-2weekly/{rangeStr}")
def load_2weekly_report(
    rangeStr: str,
    ctx = Depends(get_dept_context)
):
#-------------------------------------------------------------------------------------------------
    db, _ = ctx

    reports = db.query(TwoWeeklyReport).filter(
        TwoWeeklyReport.report_rangeStr == rangeStr
    ).order_by(asc(TwoWeeklyReport.sort_order)).all()

    return [
        {
            "uuid": r.uuid,
            "text_content": r.text_content,
            "image_uuids": [img.uuid for img in r.images]
        }
        for r in reports
    ]



#==================================================================================================
@router.get("/load-common")
def load_common_report(
    ctx = Depends(get_dept_context)
):
#--------------------------------------------------------------------------------------------------
    db, _ = ctx

    reports = db.query(CommonReport).order_by(
        asc(CommonReport.sort_order)
    ).all()

    return [
        {
            "uuid": r.uuid,
            "text_content": r.text_content
        }
        for r in reports
    ]



#==================================================================================================
@router.get("/extract-imaged/{rangeStr}")
def extract_imaged_reports(
    rangeStr: str,
    ctx = Depends(get_dept_context)
):
#--------------------------------------------------------------------------------------------------
    db, _ = ctx

    try:
        # Expected format: "2026-02-01-15"
        parts = rangeStr.split("-")
        if len(parts) != 4:
            raise ValueError("Format must be YYYY-MM-DD-DD")

        year, month, start_day, end_day = parts
        start_date = f"{year}-{month}-{start_day}"
        end_date = f"{year}-{month}-{end_day}"

        # Validates that February 30th etc. doesn't exist
        time.strptime(start_date, "%Y-%m-%d")
        time.strptime(end_date, "%Y-%m-%d")

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e) or "Invalid date format")

    reports = (
        db.query(DailyReport)
        .filter(
            and_(
                DailyReport.report_dateStr >= start_date,
                DailyReport.report_dateStr <= end_date,
                DailyReport.images.any() 
            )
        )
        .order_by(DailyReport.report_dateStr.asc(), DailyReport.sort_order.asc())
        .all()
    )

    return [
        {
            "uuid": r.uuid,
            "date": r.report_dateStr,
            "text_content": r.text_content,
            "image_uuids": [img.uuid for img in r.images]
        }
        for r in reports
    ]



#==================================================================================================
@router.get("/generate-report-docx/{rangeStr}")
def generate_report_docx(
    rangeStr: str,
    ctx = Depends(get_dept_context)
):
#--------------------------------------------------------------------------------------------------
    db, res = ctx # Get the dict
    bucket_path = res["image-bucket"] # Get the path from the dict


    try:
        # 3. Fetch the data ordered by your UI sort
        reports = db.query(TwoWeeklyReport).filter(
            TwoWeeklyReport.report_rangeStr == rangeStr
        ).order_by(TwoWeeklyReport.sort_order).all()

        if not reports:
            raise HTTPException(status_code=404, detail="No data found for this range")

        # 4. Initialize Document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        section = doc.sections[0]
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

        # Header
        doc.add_paragraph('Dear Chief Engineer')
        doc.add_paragraph('Good day Sir!,')
        doc.add_paragraph(f'Here is the 2 weekly report covering the jobs done in the period.')

        for report in reports:
            # Activity Text with List Numbering
            if report.text_content:
                para = doc.add_paragraph(report.text_content, style='List Number')
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(2)

            # Image Gallery (Auto-Flowing)
            if report.images:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                img_para.paragraph_format.space_before = Pt(0)
                img_para.paragraph_format.line_spacing = 1.15 # Slight gap between wrapped rows

                for img_record in report.images:
                    thumb_path = bucket_path / img_record.image_folder / f"{img_record.uuid}_thumb.webp"
                    
                    if thumb_path.exists():
                        with Image.open(thumb_path) as img:
                            rgb_im = img.convert('RGB')
                            img_io = io.BytesIO()
                            rgb_im.save(img_io, format='JPEG', quality=85)
                            img_io.seek(0)

                            # 1. Create a run for the picture
                            pic_run = img_para.add_run()
                            pic_run.add_picture(img_io, height=Inches(1.2))

                            # 2. Create a separate run for the spacing on the paragraph
                            img_para.add_run('   ')

        # Footer
        doc.add_paragraph('\nThank you Sir!')
        doc.add_paragraph('Best regards,')
        doc.add_paragraph('Engine Team')

        # 5. Save to Temp and Stream to Client
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        
        return FileResponse(
            temp_file.name, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"2-Weekly_Report_{rangeStr}.docx"
        )

    finally:
        db.close()

@router.get("/generate-report-docx-py/{rangeStr}")
def generate_report_docx_internal(
    rangeStr: str,
    dept_name: str = Query(None),
    date_label: str = "________",
    x_internal_key: str = Header(None)
):
    # 1. Security Check
    INTERNAL_SECRET = "Reporthub_Alpha_2026_Secure_Access"
    if x_internal_key != INTERNAL_SECRET:
        raise HTTPException(status_code=403, detail="Invalid internal key")

    # 2. Manual Context setup (Using your custom resource helper)
    resources = get_dept_resources(dept_name)
    SessionLocal = resources["sessionmaker"]
    bucket_path: Path = resources["image-bucket"]
    db = SessionLocal()

    try:
        # 3. Fetch the data ordered by your UI sort
        reports = db.query(TwoWeeklyReport).filter(
            TwoWeeklyReport.report_rangeStr == rangeStr
        ).order_by(TwoWeeklyReport.sort_order).all()

        if not reports:
            raise HTTPException(status_code=404, detail="No data found for this range")

        # 4. Initialize Document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        section = doc.sections[0]
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)

        # Header
        doc.add_paragraph('Dear Chief Engineer')
        doc.add_paragraph('Good day Sir!,')
        doc.add_paragraph(f'Here is the {date_label} covering the jobs done in the period.')

        for report in reports:
            # Activity Text with List Numbering
            if report.text_content:
                para = doc.add_paragraph(report.text_content, style='List Number')
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(2)

            # Image Gallery (Auto-Flowing)
            if report.images:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                img_para.paragraph_format.space_before = Pt(0)
                img_para.paragraph_format.line_spacing = 1.15 # Slight gap between wrapped rows

                for img_record in report.images:
                    thumb_path = bucket_path / img_record.image_folder / f"{img_record.uuid}_thumb.webp"
                    
                    if thumb_path.exists():
                        with Image.open(thumb_path) as img:
                            rgb_im = img.convert('RGB')
                            img_io = io.BytesIO()
                            rgb_im.save(img_io, format='JPEG', quality=85)
                            img_io.seek(0)

                            # 1. Create a run for the picture
                            pic_run = img_para.add_run()
                            pic_run.add_picture(img_io, height=Inches(1.2))

                            # 2. Create a separate run for the spacing on the paragraph
                            img_para.add_run('   ')

        # Footer
        doc.add_paragraph('\nThank you Sir!')
        doc.add_paragraph('Best regards,')
        doc.add_paragraph('Engine Team')

        # 5. Save to Temp and Stream to Client
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        
        return FileResponse(
            temp_file.name, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"2-Weekly_Report_{rangeStr}.docx"
        )

    finally:
        db.close()



#==================================================================================================
@router.post("/history-search")
async def start_history_search(
    query: str,
    code: str
):
#--------------------------------------------------------------------------------------------------
    # 1. Validate session
    session = SSE_SESSIONS.get(code)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    # 2. Cancel any existing search task for this code
    if code in SEARCH_TASKS:
        SEARCH_TASKS[code].cancel()
        try:
            await SEARCH_TASKS[code]
        except asyncio.CancelledError:
            pass

    # 3. Create and store the new search task
    task = asyncio.create_task(background_search_logic(query, code, session["queue"]))
    SEARCH_TASKS[code] = task
    
    return {"status": "searching"}
